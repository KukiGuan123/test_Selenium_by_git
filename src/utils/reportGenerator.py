import openpyxl
from openpyxl.drawing.image import Image
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from collections import defaultdict
from datetime import datetime
from src.utils.yaml import REPORT_PATHS, PROJECT_NAME


class ReportGenerator:
    def __init__(self):
        self.time_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.all_steps = []
        self._init_excel()
        self._init_word()

    # 初始化 Excel
    def _init_excel(self):
        # 主报表
        self.wb = openpyxl.Workbook()
        self.ws = self.wb.active
        self.ws.title = "Test Report"
        self.ws.append(["Feature", "Scenario", "Steps", "Test Result", "Error", "Screenshot"])
        for col, w in zip("ABCDEF", [16, 22, 42, 14, 20, 28]):
            self.ws.column_dimensions[col].width = w
        self.ws.row_dimensions[1].height = 10

        # 汇总报表
        self.wb_sum = openpyxl.Workbook()
        self.ws_sum = self.wb_sum.active
        self.ws_sum.title = "Summary"
        self.ws_sum.append(["Feature", "Scenario", "Test Result", "Failed Step", "Error"])
        for col, w in zip("ABCDE", [20, 35, 12, 30, 30]):
            self.ws_sum.column_dimensions[col].width = w

    # 初始化 Word
    def _init_word(self):
        self.doc = Document()
        title = self.doc.add_heading(level=1)
        run = title.add_run(f"{PROJECT_NAME}自动化测试报告")
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.table = self.doc.add_table(rows=1, cols=6)
        self.table.style = "Table Grid"
        headers = ["Feature", "Scenario", "Step", "Screenshot", "Result", "Err"]
        for i, h in enumerate(headers):
            self.table.rows[0].cells[i].text = h

        widths = [1.5, 1.5, 12, 4.0, 1.25, 1.5]
        for col, w in zip(self.table.columns, widths):
            for cell in col.cells:
                cell.width = Inches(w)


    # 添加测试步骤数据
    def add_step(self, feature, scenario, step, result, error, pic_path):
        self.all_steps.append({
            "feature": feature,
            "scenario": scenario,
            "step": step,
            "result": result,
            "error": error,
            "screenshot": pic_path
        })

    def save_all(self):
        self._save_excel()
        self._save_word()


    # 生成 Excel 主报表
    def _save_excel(self):
        for data in self.all_steps:
            row = self.ws.max_row + 1
            self.ws.row_dimensions[row].height = 90

            # 写入基础数据
            values = [data["feature"], data["scenario"], data["step"], data["result"], data["error"]]
            for i, v in enumerate(values, 1):
                cell = self.ws.cell(row=row, column=i, value=v)
                cell.alignment = Alignment(wrap_text=True, vertical="center")
                if i == 4:
                    cell.font = Font(bold=True, color="008000" if v == "Pass" else "FF0000")

            # 插入多张图片（横向排列）
            self._insert_images_in_row(data["screenshot"], row)

        # 合并主表格单元格
        self._merge_cells(self.ws, merge_cols=[1, 2])
        self.wb.save(f"{REPORT_PATHS['xls']}/{PROJECT_NAME}_详细报告_{self.time_str}.xlsx")

        # 生成 Summary
        self._build_summary()
        return f"{REPORT_PATHS['xls']}/{PROJECT_NAME}_详细报告_{self.time_str}.xlsx"

    # 【封装】批量插入图片（横向）
    def _insert_images_in_row(self, pic_path, row):
        if not pic_path:
            return

        paths = [p.strip() for p in pic_path.split(";") if p.strip() and os.path.exists(p.strip())]
        start_col = 6

        for idx, img_path in enumerate(paths):
            col = start_col + idx
            col_letter = get_column_letter(col)
            self.ws.column_dimensions[col_letter].width = 30

            try:
                img = Image(img_path)
                img.width = 160
                img.height = 120
                self.ws.add_image(img, f"{col_letter}{row}")
            except Exception as e:
                print(f"图片插入失败: {e}")

    # 【封装】生成 Summary 报表
    def _build_summary(self):
        summary = defaultdict(list)
        for s in self.all_steps:
            summary[(s["feature"], s["scenario"])].append(s)

        for (f, s), steps in summary.items():
            res = "Pass"
            fail_step = reason = ""
            for step in steps:
                if step["result"] == "Fail":
                    res = "Fail"
                    fail_step = step["step"]
                    reason = step["error"]
                    break
            self.ws_sum.append([f, s, res, fail_step, reason])
            self.ws_sum[f"C{self.ws_sum.max_row}"].font = Font(bold=True, color="008000" if res == "Pass" else "FF0000")

        # 合并 Summary
        self._merge_cells(self.ws_sum, merge_cols=[1, 2])
        self.wb_sum.save(f"{REPORT_PATHS['summary']}/{PROJECT_NAME}_详细报告_{self.time_str}.xlsx")

    # 【通用封装】合并单元格（所有表格通用）
    def _merge_cells(self, ws, merge_cols):
        if ws.max_row <= 1:
            return
        for col in merge_cols:
            self._merge_single_column(ws, col)

    # 【通用封装】单列合并（核心算法）
    def _merge_single_column(self, ws, col_idx):
        start_row = 2
        current_val = ws.cell(row=start_row, column=col_idx).value

        for row in range(start_row + 1, ws.max_row + 1):
            val = ws.cell(row=row, column=col_idx).value
            if val == current_val:
                continue
            if start_row != row - 1:
                ws.merge_cells(start_row=start_row, start_column=col_idx, end_row=row - 1, end_column=col_idx)
            start_row, current_val = row, val

        if start_row != ws.max_row:
            ws.merge_cells(start_row=start_row, start_column=col_idx, end_row=ws.max_row, end_column=col_idx)

        # 居中对齐
        for row in range(2, ws.max_row + 1):
            ws.cell(row=row, column=col_idx).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


    # 生成 Word 报告
    def _save_word(self):
        grouped = defaultdict(lambda: defaultdict(list))
        for s in self.all_steps:
            grouped[s["feature"]][s["scenario"]].append(s)

        for feature, scenarios in grouped.items():
            for scenario, steps in scenarios.items():
                for idx, step in enumerate(steps):
                    row = self.table.add_row()
                    row.height = Inches(3.0)
                    cells = row.cells
                    cells[0].text = feature if idx == 0 else ""
                    cells[1].text = scenario if idx == 0 else ""
                    cells[2].text = step["step"]
                    cells[4].text = step["result"]
                    cells[5].text = step["error"]

                    for cell in cells:
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in p.runs:
                                run.font.size = Pt(9)
                                run.font.name = "微软雅黑"

                    if step["screenshot"]:
                        try:
                            para = cells[3].paragraphs[0]
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for img in step["screenshot"].split(";"):
                                img = img.strip()
                                if os.path.exists(img):
                                    r = para.add_run()
                                    r.add_picture(img, width=Inches(1.6))
                                    para.add_run("\n")
                        except Exception as e:
                            cells[3].text = "No Image"

                    try:
                        res_run = cells[4].paragraphs[0].runs[0]
                        res_run.font.color.rgb = RGBColor(0, 128, 0) if step["result"] == "Pass" else RGBColor(255, 0,
                                                                                                               0)
                    except:
                        pass

        # 必须先合并 → 再保存！
        self._merge_word_table()

        doc_path = f"{REPORT_PATHS['doc']}/{PROJECT_NAME}_Word报告_{self.time_str}.docx"
        self.doc.save(doc_path)
        return doc_path

    # 【封装】Word 表格合并
    def _merge_word_table(self):
        merge_map = defaultdict(list)
        for i, s in enumerate(self.all_steps):
            merge_map[(s["feature"], s["scenario"])].append(i + 1)

        for key, rows in merge_map.items():
            if len(rows) > 1:
                self._merge_word_cells(self.table, 0, rows[0], rows[-1])
                self._merge_word_cells(self.table, 1, rows[0], rows[-1])


    # 【封装】Word 单元格合并
    def _merge_word_cells(self, table, col, start_row, end_row):
        try:
            table.cell(start_row, col).merge(table.cell(end_row, col))
            table.cell(start_row, col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass